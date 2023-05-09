import os
import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image
import pytesseract
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Set the path to the Tesseract executable
pytesseract.pytesseract.tesseract_cmd = r'D:/atesseract/tesseract.exe'

class ImageTextExtractor(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Image Text Extractor")
        self.geometry("700x600")
        self.configure(bg="white")

        self.image_paths = []

        self.create_widgets()

    def create_widgets(self):
        self.select_button = tk.Button(self, text="Select Images", command=self.select_images)
        self.select_button.place(relx=0.5, rely=0.2, anchor="center")

        self.image_list = tk.Listbox(self, width=50)
        self.image_list.place(relx=0.5, rely=0.4, anchor="center")

        self.delete_button = tk.Button(self, text="Delete", command=self.delete_image)
        self.delete_button.place(relx=0.5, rely=0.55, anchor="center")

        self.extract_button = tk.Button(self, text="Extract Text", command=self.extract_text)
        self.extract_button.place(relx=0.5, rely=0.7, anchor="center")

    def select_images(self):
        filetypes = [("PNG Images", "*.png")]
        selected_images = filedialog.askopenfilenames(title="Select Images", filetypes=filetypes)

        for image_path in selected_images:
            self.image_paths.append(image_path)
            self.image_list.insert(tk.END, os.path.basename(image_path))

    def delete_image(self):
        selected_indices = self.image_list.curselection()

        for index in reversed(selected_indices):
            self.image_list.delete(index)
            self.image_paths.pop(index)

    def extract_text(self):
        if not self.image_paths:
            messagebox.showerror("Error", "No images selected")
            return

        document = Document()

        for image_path in self.image_paths:
            image = Image.open(image_path)
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)

            prev_top = None
            prev_bottom = None
            paragraph = document.add_paragraph()

            for i in range(len(data["level"])):
                if data["level"][i] == 5:  # Word level
                    left, top, width, height = data["left"][i], data["top"][i], data["width"][i], data["height"][i]
                    bottom = top + height

                    if prev_bottom is not None and prev_top is not None:
                        line_spacing = top - prev_bottom

                        if line_spacing > 1.5 * (prev_bottom - prev_top):  # Detect a new paragraph based on line spacing
                            paragraph = document.add_paragraph()

                    paragraph.add_run(data["text"][i] + " ")
                    prev_top = top
                    prev_bottom = bottom

            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            document.add_paragraph()

        output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output.docx")
        document.save(output_path)
        messagebox.showinfo("Success", f"Text extracted to {output_path}")

if __name__ == "__main__":
    app = ImageTextExtractor()
    app.mainloop()

